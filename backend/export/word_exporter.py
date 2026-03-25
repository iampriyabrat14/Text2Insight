"""
Word (.docx) exporter — question, AI summary, data table, and bar chart.
"""
import io
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from backend.export.models import ExportSession
from backend.export.pdf_exporter import _build_bar_chart, _pick_chart_cols

_BLUE   = RGBColor(0x25, 0x63, 0xEB)
_INDIGO = RGBColor(0x4F, 0x46, 0xE5)
_DARK   = RGBColor(0x0F, 0x17, 0x2A)
_SLATE  = RGBColor(0x33, 0x41, 0x55)
_MUTED  = RGBColor(0x64, 0x74, 0x8B)
_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
_TH_HEX = "1E3A8A"
_ALT_HEX = "F1F5F9"


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_para_shading(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def _add_label(doc: Document, text: str, color: RGBColor):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(7.5)
    run.font.bold  = True
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    return p


def _add_data_table(doc: Document, columns: list, rows: list, max_rows: int = 30):
    """Add a styled data table to the Word document."""
    if not columns or not rows:
        return
    display = rows[:max_rows]

    tbl = doc.add_table(rows=len(display) + 1, cols=len(columns))
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for j, col in enumerate(columns):
        cell = hdr_cells[j]
        cell.text = col[:24]
        run = cell.paragraphs[0].runs[0]
        run.font.bold      = True
        run.font.size      = Pt(8)
        run.font.color.rgb = _WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, _TH_HEX)

    # Data rows
    for i, row in enumerate(display, start=1):
        cells = tbl.rows[i].cells
        bg    = _ALT_HEX if i % 2 == 0 else "FFFFFF"
        for j, col in enumerate(columns):
            val = row.get(col)
            s   = str(val) if val is not None else "—"
            cell = cells[j]
            cell.text = s[:40]
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8)
            _set_cell_bg(cell, bg)

    doc.add_paragraph()  # breathing space after table


def _add_chart_image(doc: Document, columns: list, rows: list):
    """Render bar chart via ReportLab renderPM and embed as picture. Best-effort."""
    drawing = _build_bar_chart(columns, rows)
    if drawing is None:
        return
    try:
        from reportlab.graphics import renderPM
        png = renderPM.drawToString(drawing, fmt="PNG", dpi=130)
        doc.add_picture(io.BytesIO(png), width=Cm(14))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass   # silently skip if renderPM / Pillow not available


def generate_word(session: ExportSession) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    # ── Cover ─────────────────────────────────────────────────────────────
    title_p = doc.add_heading("AI Insights Report", level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.color.rgb = _BLUE
        run.font.size      = Pt(26)

    sub = doc.add_paragraph(session.title)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size      = Pt(13)
        run.font.color.rgb = _SLATE

    export_dt = datetime.now(timezone.utc).strftime("%B %d, %Y")
    meta = doc.add_paragraph(
        f"Prepared by {session.username}  ·  {export_dt}  ·  {len(session.qa_pairs)} questions"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size      = Pt(9)
        run.font.color.rgb = _MUTED

    doc.add_page_break()

    # ── Q&A pairs ─────────────────────────────────────────────────────────
    for idx, pair in enumerate(session.qa_pairs, start=1):
        # Question
        _add_label(doc, f"QUESTION {idx}", _BLUE)
        q_para = doc.add_paragraph()
        q_run  = q_para.add_run(pair["query"])
        q_run.font.size  = Pt(11)
        q_run.font.bold  = True
        q_run.font.color.rgb = _DARK
        q_para.paragraph_format.left_indent  = Cm(0.5)
        q_para.paragraph_format.right_indent = Cm(0.5)
        q_para.paragraph_format.space_after  = Pt(8)
        _set_para_shading(q_para, "EFF6FF")

        # Summary
        _add_label(doc, "AI SUMMARY", _INDIGO)
        a_para = doc.add_paragraph(pair["summary"])
        for run in a_para.runs:
            run.font.size      = Pt(11)
            run.font.color.rgb = _SLATE
        a_para.paragraph_format.space_after = Pt(10)

        # Data table + chart
        result  = pair.get("result_data") or {}
        columns = result.get("columns", [])
        rows    = result.get("rows",    [])

        if columns and rows:
            _add_label(doc, "DATA TABLE", _MUTED)
            _add_data_table(doc, columns, rows)

            _add_label(doc, "CHART", _MUTED)
            _add_chart_image(doc, columns, rows)

        # Divider
        hr = doc.add_paragraph()
        hr.paragraph_format.space_before = Pt(4)
        hr.paragraph_format.space_after  = Pt(4)
        pBdr   = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "E2E8F0")
        pBdr.append(bottom)
        hr._p.get_or_add_pPr().append(pBdr)

    if not session.qa_pairs:
        doc.add_paragraph("No queries in this session.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
